import json
from docx import Document
from datetime import datetime, timezone, timedelta
from docx.shared import Cm

class WordConverter:
    def __init__(self, module_name:str, test_name:str, report_data:json) -> None:
        self.document = Document()
        self.report_data = report_data
        self.module_name = module_name
        self.test_name = test_name

    def create_word_report(self):
        doc = self.document

        # 添加標題
        doc.add_heading(f'Report for {self.module_name}: {self.test_name}', 0)

        #從報告中再提取資訊
        collection = self.report_data['collection']['info']['name']
        run_info = self.report_data['run']

        # 總結
        main_info = doc.add_paragraph(f'Collection: {collection}\n')
        time = datetime.fromtimestamp(run_info['timings']['completed']/1000).astimezone(timezone(timedelta(hours=8))).strftime('%Y-%m-%d %H:%M:%S GMT+8')
        main_info.add_run(f'Time: {time}\n')

        #表1
        self.add_table_1()

        total_info = doc.add_paragraph()

        #測試時長
        time_duration = (run_info['timings']['completed']-run_info['timings']['started'])/1000
        time_text = self.time_switch(time_duration)
        total_info.add_run(f'Total run duration: {time_text}\n')

        #總接收資料
        total_data_received = run_info['transfers']['responseTotal']
        tdr_str = self.datasize_switch(total_data_received)
        total_info.add_run(f'Total data received: {tdr_str}\n')

        #平均回應時間
        response_time = run_info['timings']['responseAverage']
        total_info.add_run(f'Average response time: {response_time:.2f} ms')

        #總失敗次數
        total_fail = doc.add_paragraph().add_run('Total Failures: ')
        if run_info['failures']:
            total_fail.add_text(str(len(run_info['failures'])))
        else:
            total_fail.add_text('0')
        total_fail.bold = True

        #測試詳細結果
        self.add_request_details()

        # 保存 Word 文件
        try:
            doc.save(f'docx/report_{self.module_name}_{self.test_name}.docx')
            print(f"Document saved as report_{self.module_name}_{self.test_name}.docx")
        except Exception as e:
            print(f"Error saving document: {e}")

    def add_table_1(self):
        doc = self.document
        run_info = self.report_data['run']

        summary_table = doc.add_table(6, 3)
        summary_table.style = 'Light Shading Accent 5'           #設定表格樣式
        summary_table.cell(0, 1).text = 'Total'         #col欄位名稱
        summary_table.cell(0, 2).text = 'Failed'

        rows = ['Iterations','Requests','Prerequest Scripts', 'Test Scripts', 'Assertions']             #row欄位名稱
        rows_json_name = ['iterations','requests','prerequestScripts', 'testScripts', 'assertions']
        for i in range(len(rows)):
            summary_table.cell(i+1, 0).text = rows[i]
            summary_table.cell(i+1, 1).text = str(run_info['stats'][rows_json_name[i]]['total'])
            summary_table.cell(i+1, 2).text = str(run_info['stats'][rows_json_name[i]]['failed'])

    # def add_table_2(self):
    #     doc = self.document
    #     run_info = self.report_data['run']

    #     summary_table2 = doc.add_table(3, 2)
    #     summary_table2.style = 'Light Grid'             #設定表格樣式
    #     rows_2 = ['Total run duration', 'Total data received', 'Average response time']              #row欄位名稱
    #     for i in range(len(rows_2)):
    #         summary_table2.cell(i, 0).text = rows_2[i]

    #     time_duration = datetime.fromtimestamp((run_info['timings']['completed']-run_info['timings']['started'])/1000)
    #     summary_table2.cell(0, 1).text = time_duration.strftime('%Mm %Ss')

    #     total_data_received = run_info['transfers']['responseTotal']
    #     if(total_data_received>1024):
    #         if(total_data_received>1024*1024):
    #             tdr_str = f"{total_data_received/(1024*1024):.2f}" + 'MB'
    #         else:
    #             tdr_str = f"{total_data_received/1024:.2f}" + 'KB'
    #     else:
    #         tdr_str = f"{total_data_received:.2f}" + 'B'
    #     summary_table2.cell(1, 1).text = tdr_str

    #     summary_table2.cell(2, 1).text = str(run_info['timings']['responseAverage'])+'ms'

    def time_switch(self, time):
        date_time = datetime.fromtimestamp(time)

        if time>60*60:
            return str(date_time.strftime('%Hh %Mm %Ss'))
        elif time>60:
            return str(date_time.strftime('%Mm %Ss'))
        else:
            return str(date_time.strftime('%Ss'))

    def datasize_switch(self, size):
        if(size>1024*1024):
            return f"{size/(1024*1024):.2f} MB"
        elif(size>1024):
            return f"{size/1024:.2f} KB"
        else:
            return f"{size:.2f} B"

    def add_request_details(self):
        #requests結果
        doc = self.document
        doc.add_heading('Requests', 1 )

        execution_info = self.report_data['run']['executions']
        for i, exec_info in enumerate(execution_info):
            name = exec_info['item']['name']
            doc.add_heading(name, 2)

            #request detail
            request_info = doc.add_paragraph()
            request_info.add_run('Method: '+ exec_info['request']['method'] +'\n')
            url = self.make_url(exec_info['request']['url'])
            request_info.add_run(f'URL: {url}\n')
            request_info.add_run('Status code: ' + str(exec_info['response']['code']))

            #response detail
            response_info = doc.add_paragraph()
            response_info.add_run('Response Time: ' + str(exec_info['response']['responseTime']) + 'ms\n')
            response_info.add_run('Response Size: ' + str(exec_info['response']['responseSize']) + 'B\n')


            doc.add_heading('Tests', 3)
            test_num = len(exec_info['assertions'])
            tests_table = doc.add_table(test_num+1, 3)
            tests_table.style = 'Light Shading Accent 5'
            tests_table_cols = ['Name', 'Pass count', 'Fail count']
            for j in range (len(tests_table_cols)):
                tests_table.cell(0, j).text = tests_table_cols[j]
                # tests_table.cell(0, j).add_paragraph().add_run(tests_table_cols[j]).bold = True
            assertions = exec_info['assertions']
            for j, assert_info in enumerate(assertions):
                tests_table.cell(j+1, 0).text = assert_info['assertion']
                if 'error' in assert_info:
                    tests_table.cell(j+1, 2).text = '1'
                    tests_table.cell(j+1, 1).text = '0'
                    self.add_failure_detail(assert_info['error'])
                else:
                    tests_table.cell(j+1, 2).text = '0'
                    tests_table.cell(j+1, 1).text = '1'
            doc.add_paragraph('\n')

    def add_failure_detail(self, error_info):
        doc = self.document
        failure = False
        if not failure:
            doc.add_heading('Failure', 3)
            failure = True

        fail_table = doc.add_table(1,2)
        fail_table.style ='Light Shading Accent 2'
        fail_table.allow_autofit = False
        fail_table.columns[0].width = Cm(5)
        fail_table.columns[1].width = Cm(11)
        # for cell in fail_table.columns[0].cells:
        #     cell.width = Cm(4)
            # fail_table.rows[0].cells[0].width = Cm(6)
        fail_table.cell(0,0).text = error_info['test']
        fail_table.cell(0,1).text = error_info['message']


    def make_url(self, url_data):
        protocol = url_data['protocol']

        host = ''
        for host_e in url_data['host']:
            host += host_e + '.'
        host = host.rstrip('.')  #移除最後多餘的點

        port = url_data.get('port', '')

        path = ''
        for path_e in url_data['path']:
            path += '/' + path_e

        query = ''
        if(url_data['query']):
            query += '?'
            for query_e in url_data['query']:
                query += query_e['key'] + '=' + query_e['value'] + '&'
            query = query.rstrip('&')

        # 組合完整 URL
        if port:
            return f'{protocol}://{host}:{port}{path}{query}'
        else:
            return f'{protocol}://{host}{path}{query}'

